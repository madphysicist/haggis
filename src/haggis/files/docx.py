#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# haggis: a library of general purpose utilities
#
# Copyright (C) 2019  Joseph R. Fox-Rabinovitz <jfoxrabinovitz at gmail dot com>
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# Author: Joseph Fox-Rabinovitz <jfoxrabinovitz at gmail dot com>
# Version: 13 Apr 2019: Initial Coding


"""
Utilities for working with new-style MS Word documents, only available
when the ``[docx]`` :ref:`extra <installation-extras>` is installed.

If `python-docx`_ is not found at import time, this module will have a
:py:data:`docx_enabled` attribute, which will be :py:obj:`False`. If
`python-docx`_ is found, on the other hand, :py:data:`docx_enabled` will
be :py:obj:`True`, and all the functions and attributes of the module
will be present.

Some of the methods here are workarounds for features that are missing
or buggy in the original library. Often, these methods are inspired by
recipes found in the corresponding bug reports and Stack Overflow posts,
which are referenced as appropriate.

.. py:data:: docx_enabled

   A boolean value indicating whether the ``[docx]``
   :ref:`extra <installation-extras>` has been installed. If
   :py:obj:`False`, the API will be severely limited.


.. include:: /link-defs.rst 
"""


__all__ = ['docx_enabled',]


from collections import OrderedDict
from os.path import dirname, join
import re

try:
    from docx.document import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.shared import RGBColor, Inches, Cm, Mm, Pt, Twips, Emu
    from docx.table import _Cell, Table
    from docx.text.run import Run
    from docx.text.paragraph import Paragraph

    from lxml import etree
except ImportError:
    from .. import _display_missing_extra
    _display_missing_extra('docx', ['python-docx', 'lxml'])
    docx_enabled = False
else:
    __all__.extend([
        'block_iterator', 'insert_toc', 'merge_row', 'set_row_height',
        'style_row', 'style_column', 'table_no_fill', 'edit_font',
        'is_paragraph_empty', 'delete_paragraph', 'list_number', 'add_section',
        'insert_math_ml', 'str2length',
    ])
    docx_enabled = True


if docx_enabled:
    def block_iterator(parent):
        """
        Yield each paragraph and table child within `parent`, in
        document order.

        Each returned value is an instance of either
        :py:class:`~docx.table.Table` or
        :py:class:`~docx.text.paragraph.Paragraph`. `parent` would most
        commonly be a reference to a main
        :py:class:`~docx.document.Document` object, but also works for a
        :py:class:`~docx.table._Cell` object, which itself can contain
        paragraphs and tables.

        This function is taken verbatim from python-docx
        `Issue #40 <https://github.com/python-openxml/python-docx/issues/40>`_.
        Hopefully it will make it to `python-docx`_ soon.


        .. include:: /link-defs.rst 
        """
        if isinstance(parent, Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)


    def insert_toc(doc, min_level=1, max_level=3):
        """
        Insert a table of contents stub into a
        :py:class:`~docx.document.Document`.

        The TOC is a stub and needs to be updated manually when the
        end-user opens the generated document in a Word client.

        This function is taken almost verbatim from @mustash's comment
        to python-docx `Issue #36 <https://github.com/python-openxml/python-docx/issues/36#issuecomment-145302669>`_.
        See the previous comment in the thread for more information.
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement('w:fldChar')          # creates a new element
        fld_char.set(qn('w:fldCharType'), 'begin')   # sets attribute on element
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instr_text.text = 'TOC \\o "{min}-{max}" \\h \\z \\u'.format(
                            min=int(min_level), max=int(max_level))

        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')
        fld_char3 = OxmlElement('w:t')
        fld_char3.text = 'Right-click to update field.'
        fld_char2.append(fld_char3)

        fld_char4 = OxmlElement('w:fldChar')
        fld_char4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fld_char)
        r_element.append(instr_text)
        r_element.append(fld_char2)
        r_element.append(fld_char4)


    def merge_row(table, row=0):
        """
        Merge a row in a :py:class:`~docx.table.Table` into a single
        cell.

        Operates on the first row by default.
        """
        rc = table.row_cells(row)
        rc[0].merge(rc[-1])


    def set_row_height(row, height):
        """
        Set the height of a :py:class:`docx.table.Table` row.

        Parameters
        ----------
        row : docx.table._Row
            The row index.
        height : docx.shared.Length
            The height to assign.


        This function will be obsolted by python-docx
        `Pull Request #301 <https://github.com/python-openxml/python-docx/pull/301>`_.
        """
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height.twips))

        trPr = OxmlElement('w:trPr')
        trPr.append(trHeight)

        row._tr.append(trPr)


    def style_row(doc, style, *args):
        """
        .. py:function:: style_row(doc, style, row)
        .. py:function:: style_row(doc, style, table, rowid)

        Set a custom style for all the text in a row in the table.

        Every paragraph of each cell in the row will be styled, so use
        with care.

        Parameters
        ----------
        doc : docx.document.Document
            The document containing the items to style. Only required 
            when `style` is set by name.
        style : docx.styles.style._CharacterStyle or docx.styles.style._ParagraphStyle or str
            The name or style object representing a run- or
            paragraph-level style to apply.
        row : docx.table._Row
            The row object to style.
        table : docx.table.Table
            The table to style.
        rowid : int
            The index of the row to style.
        """
        if len(args) == 1:
            row, = args
        elif len(args) == 2:
            table, rowid = args
            row = table.rows[rowid]
        else:
            raise TypeError(
                'Expected 1 or 2 positional arguments, got {}'.format(len(args))
            )

        if isinstance(style, str):
            style = doc.styles[style]

        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            def apply(par, style):
                par.style = style
        elif style.type == WD_STYLE_TYPE.CHARACTER:
            def apply(par, style):
                for run in par.runs:
                    run.style = style
        else:
            raise ValueError('Paragraph and character styles only.')

        for cell in row.cells:
            for par in cell.paragraphs:
                apply(par, style)


    def style_column(doc, style, *args):
        """
        .. py:function:: style_column(col, *, style)
        .. py:function:: style_column(table, colid, *, style)

        Set a custom style for all the text in a column in the table.

        Every paragraph of each cell in the column will be styled, so
        use with care.

        Parameters
        ----------
        doc : docx.document.Document
            The document containing the items to style. Only required
            when `style` is set by name.
        style : docx.styles.style._CharacterStyle or docx.styles.style._ParagraphStyle or str
            The name or style object representing a run- or
            paragraph-level style to apply.
        col : docx.table._Column
            The column object to style.
        table : docx.table.Table
            The table to style.
        colid : int
            The index of the column to style.
        """
        if len(args) == 1:
            col, = args
        elif len(args) == 2:
            table, colid = args
            col = table.columns[colid]
        else:
            raise TypeError(
                'Expected 1 or 2 positional arguments, got {}'.format(len(args))
            )

        if isinstance(style, str):
            style = doc.styles[style]

        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            def apply(par, style):
                par.style = style
        elif style.type == WD_STYLE_TYPE.CHARACTER:
            def apply(par, style):
                for run in par.runs:
                    run.style = style
        else:
            raise ValueError('Paragraph and character styles only.')

        for cell in col.cells:
            for par in cell.paragraphs:
                apply(par, style)


    def table_no_fill(table):
        """
        Set the table not to fill up the entire page width.

        This should be called after the table has been filled in, or at
        least the number of rows and columns has been set. Adding rows
        and columns may cause unpredictable layout conflicts which may
        invalidate these settings.

        This method is a result of research done on
        https://github.com/python-openxml/python-docx/issues/315.
        """
        for row in table.rows:
            for col in row._tr.tc_lst:
                tcW = col.tcPr.tcW
                tcW.type = 'auto'
                tcW.w = 0


    #: Mapping of supported :py:class:`~docx.text.run.Font` attributes
    #: to dictionaries describing how to set them. All nested keys are
    #: optional. Options are:
    #:
    #:   - ``parser``: Method that converts inputs to values
    #:     `python-docx`_ understands.
    #:   - ``formatter``: Method that converts inputs from internal
    #:     `python-docx`_ values to user-friendly ones.
    #:   - ``setter``: Alternative method to set the attribute on the
    #:     :py:class:`~docx.text.run.Font` object.
    #:
    #: Used by :func:`edit_font`.
    _known_font_attributes = OrderedDict([
        ('name', {'formatter': lambda x: ' + "{}"'.format(x)}),
        ('size', {
            'parser': lambda x: Pt(x),
            'formatter': lambda x: ' + {}pt'.format(x.pt),
        }),
        ('color', {
            'parser': lambda x: RGBColor(*x),
            'formatter': lambda x: ' + rgb({}, {}, {})'.format(*x),
            'setter': lambda font, attr, x: setattr(font.color, 'rgb', x)
        }),
        ('all_caps', {}),
        ('bold', {}),
        ('double_strike', {}),
        ('emboss', {}),
        ('imprint', {}),
        ('italic', {}),
        ('outline', {}),
        ('shadow', {}),
        ('small_caps', {}),
        ('strike', {}),
        ('subscript', {}),
        ('superscript', {}),
        ('underline', {}),
    ])


    def edit_font(doc, runs, **kwargs):
        """
        Adds a dynamic run-style to the document that changes the font
        properties of the specified runs.

        The updated style is generated on the fly and inherits all
        properties except the ones that are modified from the current
        character style of the run. The style name is the original style
        name with the modified properties appended after a ``+`` or
        ``-`` sign, in alphabetical order. If such a character style
        already exists in the document, it will be used instead of
        making a new one.

        The `runs` parameter is expected to generally come from the
        :py:attr:`~docx.text.paragraph.Paragraph.runs` attribute of a
        :py:class:`~docx.text.paragraph.Paragraph` object. A single
        :py:class:`~docx.text.run.Run` can be processed as well.

        Parameters
        ----------
        doc : docx.document.Document
            The document containing the runs to modify. This is
            necessary because the styles must be added to the document.
        runs : iterable[docx.text.run.Run] or docx.text.run.Run
            Either an iterable or runs or a single run. The runs will be
            given a new character style that inherits from the current
            one but contains an updated font with the requested color.
        name: str
            The name of the font. :py:obj:`None` if not to be modified.
            Added as ``... + "NAME"``.
        size: float or int
            The font size in points. :py:obj:`None` if not to be
            modified. Added as ``... + SIZEpt``.
        color: tuple(int, int, int)
            A three-element tuple of integers that represents the RGB
            components of the color to set, or `None` if color is not to
            be modified. Other iterables are accepted, but they are
            converted to tuples internally. Added as
            ``... + rgb(R,G,B)``.


        All remaining arguments are processed as tri-state booleans,
        where :py:obj:`None` indicates no modification. Flags are
        appended as ``... + FLAG`` if :py:obj:`True`, ``... - FLAG`` if
        :py:obj:`False`. Supported arguments flags are: 

          - `bold`
          - `italic`
          - `underline`
          - `subscript`
          - `superscript`
          - `all_caps`
          - `emboss`
          - `strike`
          - `double_strike`
          - `imprint`
          - `outline`
          - `shadow`
          - `small_caps`
        """
        if runs is None:
            return

        if isinstance(runs, Run):
            return edit_font(doc, [runs], **kwargs)

        for run in runs:
            old_style = run.style
            if old_style is None:
                old_style = doc.styles.default(WD_STYLE_TYPE.CHARACTER)
            new_name = old_style.name
            for attr in kwargs:
                value = kwargs[attr]
                if value is not None:
                    properties = _known_font_attributes[attr]
                    if 'parser' in properties:
                        value = properties['parser'](value)
                        kwargs[attr] = value
                    if 'formatter' in properties:
                        new_name += properties['formatter'](value)
                    else:
                        new_name += ' {} {}'.format('+' if value else '-', attr)
            try:
                new_style = doc.styles[new_name]
            except KeyError:
                new_style = doc.styles.add_style(new_name, WD_STYLE_TYPE.CHARACTER)
            new_style.base_style = old_style
            font = new_style.font
            for attr, value in kwargs.items():
                properties = _known_font_attributes[attr]
                if 'setter' in properties:
                    properties['setter'](font, attr, value)
                else:
                    setattr(font, attr, value)
            run.style = new_style


    def is_paragraph_empty(paragraph):
        """
        Check if a paragraph contains content.

        Content includes any item besides styling information and empty runs.

        Source: https://stackoverflow.com/q/51217113/2988730
        """
        p = paragraph._p
        runs = p.xpath('./w:r[./*[not(self::w:rPr)]]')
        others = p.xpath('./*[not(self::w:pPr) and not(self::w:r) and '
                         'not(contains(local-name(), "bookmark"))]')
        return len(runs) + len(others) == 0


    def delete_paragraph(paragraph):
        """
        Remove a paragraph from an enclosing document.

        This may or may not work properly if the paragraph contains
        non-trivial content, like pictures.

        Inspiration is from @scanny's comment regarging python-docx
        `Issue #33 <https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907>`_.
        """
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None


    def list_number(doc, par, prev=None, level=None, num=True):
        """
        Make a paragraph into a list item with a specific level and
        optional restart.

        An attempt will be made to retreive an abstract numbering style
        that corresponds to the style of the paragraph. If that is not
        possible, the default numbering or bullet style will be used
        based on the `num` parameter.

        Parameters
        ----------
        doc : ~docx.document.Document
            The document to add the list into.
        par : ~docx.text.paragraph.Paragraph
            The paragraph to turn into a list item.
        prev : ~docx.text.paragraph.Paragraph or None
            The previous paragraph in the list. If specified, the
            numbering and styles will be taken as a continuation of this
            paragraph. If omitted, a new numbering scheme will be
            started.
        level : int or None
            The level of the paragraph within the outline. If `prev` is
            set, defaults to the same level as in `prev`. Otherwise,
            defaults to zero.
        num : bool
            If `prev` is :py:obj:`None` and the style of the paragraph
            does not correspond to an existing numbering style, this
            will determine wether or not the list will be numbered or
            bulleted. The result is not guaranteed, but is fairly safe
            for most Word templates.


        The code here is mainly taken from python-docx
        `Issue #25 <https://github.com/python-openxml/python-docx/issues/25>`_
        and `Pull Request #110 <https://github.com/python-openxml/python-docx/pull/110>`_
        In particular, see the two comments by @yurac:
        `[1] <https://github.com/python-openxml/python-docx/issues/25#issuecomment-140334543>`_
        and `[2] <https://github.com/python-openxml/python-docx/issues/25#issuecomment-143231954>`_
        """
        xpath_options = {
            True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
            False: {'single': '', 'level': level},
        }

        def style_xpath(prefer_single=True):
            """
            The style comes from the outer-scope variable ``par.style.name``.
            """
            style = par.style.style_id
            return (
                'w:abstractNum['
                    '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
                ']/@w:abstractNumId'
            ).format(style=style, **xpath_options[prefer_single])

        def type_xpath(prefer_single=True):
            """
            The type is from the outer-scope variable ``num``.
            """
            type = 'decimal' if num else 'bullet'
            return (
                'w:abstractNum['
                    '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
                ']/@w:abstractNumId'
            ).format(type=type, **xpath_options[prefer_single])

        def get_abstract_id():
            """
            Select as follows:

                1. Match single-level by style (get min ID)
                2. Match exact style and level (get min ID)
                3. Match single-level decimal/bullet types (get min ID)
                4. Match decimal/bullet in requested level (get min ID)
                3. 0
            """
            for fn in (style_xpath, type_xpath):
                for prefer_single in (True, False):
                    xpath = fn(prefer_single)
                    ids = numbering.xpath(xpath)
                    if ids:
                        return min(int(x) for x in ids)
            return 0

        if (prev is None or
                prev._p.pPr is None or
                prev._p.pPr.numPr is None or
                prev._p.pPr.numPr.numId is None):
            if level is None:
                level = 0
            numbering = doc.part.numbering_part.numbering_definitions._numbering
            # Compute the abstract ID first by style, then by num
            anum = get_abstract_id()
            # Set the concrete numbering based on the abstract numbering ID
            num = numbering.add_num(anum)
            # Make sure to override the abstract continuation property
            num.add_lvlOverride(ilvl=level).add_startOverride(1)
            # Extract the newly-allocated concrete numbering ID
            num = num.numId
            ret = num, anum

        else:
            # Get the previous concrete numbering ID
            if level is None:
                level = prev._p.pPr.numPr.ilvl.val
            num = prev._p.pPr.numPr.numId.val
            ret = num
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level
        return ret


    def add_section(doc, orientation=WD_ORIENTATION.PORTRAIT):
        """
        Add a new section to `doc` with the specified page orientation.

        This function always creates a new section with page break style
        :ref:`wdsectionstart`.NEW_PAGE, even if
        the previous section has the same orientation. The width and
        height of the new section will be swapped if necessary so that
        the width is greater in landscape mode and the height is greater
        in portrait mode.

        `orientation` may be one of the :ref:`wdorientation` enums, or
        the strings {``'portrait'``, ``'landscape'``} (case
        insensitive).

        Returns the newly created section.

        This function is a workaround for python-docx bug
        `#214 <https://github.com/python-openxml/python-docx/issues/214>`_.
        The same workaround is also described in
        http://stackoverflow.com/q/31893557/2988730.
        """
        if isinstance(orientation, str):
            orientation = orientation.lower()
            if orientation == 'portrait':
                orientation = WD_ORIENTATION.PORTRAIT
            elif orientation == 'landscape':
                orientation = WD_ORIENTATION.LANDSCAPE
            else:
                raise ValueError(
                        'Invalid page orientation "{}"'.format(orientation)
                )
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section.orientation = orientation
        width, height = section.page_width, section.page_height
        if (orientation == WD_ORIENTATION.PORTRAIT and width > height) or \
           (orientation == WD_ORIENTATION.LANDSCAPE and height > width):
            section.page_width, section.page_height = height, width
        return section


    try:
        # realpath/abspath omitted intentionally
        _mml2omml_file = join(dirname(__file__), '_resources', 'MML2OMML.XSL')
        _mml2omml_xslt = etree.parse(_mml2omml_file)
    except (OSError, FileNotFoundError):
        def insert_math_ml(*args, **kwargs):
            raise NotImplementedError(
                'Unable to find XSL file at "{}"'.format(_mml2omml_file)
            )
    else:
        _mml2omml_transform = etree.XSLT(_mml2omml_xslt)
        def insert_math_ml(par, math_ml):
            """
            Convert a MathML equation to an Open MathML format suitable
            for MS documents, and insert it into the specified
            paragraph.

            The MathML is converted to Open MathML format using an
            internal stylesheet and inserted into the end of the
            paragraph.

            Parameters
            ----------
            par : ~docx.text.paragraph.Paragraph
                The paragraph to append the equation to.
            math_ml : str or file-like
                If a string that starts with an opening ``<math>`` tag
                and ends with a closing ``</math>`` tag, it will be
                parsed as MathML. Other wise it will be interpreted as a
                file name or file object, as appropriate.

            Notes
            -----
            If the MML to OMML stylesheet can not be located for any
            reason, this function will raise a
            :py:exc:`NotImplementedError`. The error will indicate the
            expected location of the file.

            This function and its setup are based on the discussion of
            `issue #320 <https://github.com/python-openxml/python-docx/issues/320>`_
            on the python-docx GitHub page. The file MML2OMML.XSL was provided
            by user @peepall as part of the discussion.
            """
            loader = etree.parse
            if isinstance(math_ml, str):
                mmc = math_ml.strip().casefold()
                if mmc.startswith('<math') and mmc.endswith('</math>'):
                    loader = etree.fromstring

            mml = loader(math_ml)
            omml = _mml2omml_transform(mml)
            par._element.append(omml.getroot())


    #: Mapping of recognized unit names to their corresponding
    #: `python-docx`_ classes.
    _units_mapping = {
        'emu': Emu,
        'english metric unit': Emu,
        'english metric units': Emu,
        'englishmetricunit': Emu,
        'englishmetricunits': Emu,
        'cm': Cm,
        'centimeter': Cm,
        'centimeters': Cm,
        '"': Inches,
        'in': Inches,
        'inch': Inches,
        'inches': Inches,
        'millimeter': Mm,
        'millimeters': Mm,
        'mm': Mm,
        'pt': Pt,
        'point': Pt,
        'points': Pt,
        'twip': Twips,
        'twips': Twips,
    }


    #: Regular expression to extract floats from a string.
    _float_pattern = re.compile(
        r'[+-]?(?:(?:\d*\.\d+)|(?:\d+\.?))(?:[Ee][+-]?\d+)?'
    )


    def str2length(string, default_units='"'):
        """
        Convert a string with a floating point number and associated
        units into a :py:class:`docx.shared.Length`.

        The number and the units may be separated by optional
        whitespace. If the units are missing entirely, `default_units`
        will be used instead. `default_units` may be a
        :py:class:`docx.shared.Length` subclass, or a string naming the
        units. The default default is :py:class:`docx.shared.Inches`.

        Unit names are case insensitive. They may be written out fully or
        abbreviated, plural or singular.
        """
        string = str(string)
        match = _float_pattern.match(string)
        if match is None:
            raise ValueError('Must begin with valid floating point component')
        num = float(match.group())
        units = string[match.end():].strip()

        if not units:
            units = default_units

        if isinstance(units, str):
            cunits = units.casefold()
            if cunits not in _units_mapping:
                raise ValueError('Unknown units {!r}'.format(units))
            units = _units_mapping[cunits]

        return units(num)
